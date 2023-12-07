import pyodbc
import csv
from datetime import datetime, timedelta
from docx import Document
from collections import defaultdict

# Database connection parameters
server = "soda.aapm.org"
database = "nfaapm"
username = "sa"
password = "LP7ka4M"
cnxn = pyodbc.connect(f"DRIVER={{ODBC Driver 17 for SQL Server}};SERVER={server};DATABASE={database};UID={username};PWD={password}")

# Create a cursor object
cursor = cnxn.cursor()

# Day and time ranges
days = {"Saturday": {"start": datetime(2023, 12, 9, 8, 0), "end": datetime(2023, 12, 9, 22, 0)},
        "Sunday": {"start": datetime(2023, 12, 10, 8, 0), "end": datetime(2023, 12, 10, 19, 0)}}

# Number of concurrent rooms
num_rooms = 28

# Get committees with duration
committee_query = """
SELECT evt_id, cmt_code, que_response FROM [RoomScheduler].[dbo].[SurveyResponse]
WHERE que_id = 4 AND evt_id = 38"""
cursor.execute(committee_query)
committees = []
for row in cursor.fetchall():
    committees.append((row[1], row[2]))


# Define a function to convert duration string to minutes
def convert_duration(duration_str):
    # Assume duration_str is something like '1 hour' or '30 minutes'
    # This function should parse the duration_str and return a timedelta object
    try:
        parts = duration_str.split()
        if 'hour' in duration_str or 'hours' in duration_str:
            hours = int(parts[0])
            return timedelta(hours=hours)
        elif 'minute' in duration_str or 'minutes' in duration_str:
            minutes = int(parts[0])
            return timedelta(minutes=minutes)
    except (ValueError, IndexError):
        print(f"Invalid duration format: {duration_str}")
        return timedelta()




def has_conflict(new_meeting, scheduled_meetings):
    """
    Checks if the new_meeting conflicts with any of the scheduled_meetings.
    A conflict occurs if a meeting is scheduled in the same room at an overlapping time.

    :param new_meeting: A dictionary representing the meeting to be scheduled.
    :param scheduled_meetings: A list of dictionaries representing already scheduled meetings.
    :return: True if there is a conflict, False otherwise.
    """
    for scheduled_meeting in scheduled_meetings:
        # Check for room conflict
        if new_meeting["room"] == scheduled_meeting["room"]:
            # Check for time overlap
            new_start = new_meeting["start_time"]
            new_end = new_meeting["end_time"]
            scheduled_start = scheduled_meeting["start_time"]
            scheduled_end = scheduled_meeting["end_time"]

            # Overlap occurs if the new meeting starts before the scheduled one ends,
            # and the new meeting ends after the scheduled one starts
            if new_start < scheduled_end and new_end > scheduled_start:
                return True

    # No conflicts found
    return False



# Get attendees for each committee
attendees_dict = defaultdict(list)
for committee in committees:
    committee_code, _ = committee
    committee_member_query = f"""
    select ind_id, 'Y' as [ischair], committee_code, 'Y' as [voting], ind_first_name, ind_last_name, email from [dbo].[AAPM_vwcommitteechairs] where committee_code = '{committee_code}'
    union
    SELECT
        Ind_id, 'N' as ischair,
        Committee_Code,
        MAX(CASE WHEN Voting = 'Y' THEN 'Y' ELSE 'N' END) as Voting, ind_first_name, ind_last_name, email
    FROM
        AAPM_vwCommitteeMembership
    WHERE
        Committee_Code = '{committee_code}'
        AND Function_Code NOT LIKE '%GST'
        and ind_id not in (select ind_id from [dbo].[AAPM_vwcommitteechairs] where committee_code = '{committee_code}')
    GROUP BY
        Ind_id,
        Committee_Code, ind_first_name, ind_last_name, email
    HAVING
        MAX(CASE WHEN Voting = 'Y' THEN 'Y' ELSE 'N' END) = 'Y'
    ORDER BY
        Ind_id;
    """
    cursor.execute(committee_member_query)
    for row in cursor.fetchall():
        attendees_dict[committee_code].append(
            {
                "ind_id": row[0],
                "ischair": row[1],
                "voting": row[3],
                "name": f"{row[4]} {row[5]}",
                "email": row[6],
            }
        )

# Define a function to check for conflicts with other attendees
def has_attendee_conflict(meeting, scheduled_meetings):
    for scheduled_meeting in scheduled_meetings:
        if meeting["committee"] != scheduled_meeting["committee"]:
            continue
        for attendee in meeting["attendees"]:
            if attendee in scheduled_meeting["attendees"]:
                return True
    return False

# Define a function to schedule a meeting
def schedule_meeting(committee, duration, attendees, scheduled_meetings):
    print(f"Attempting to schedule meeting for committee: {committee} with duration {duration}")
    for day in days.keys():
        print(f"Checking availability on {day}")
        start_time = days[day]["start"]
        end_time = days[day]["end"]
        meeting_duration = convert_duration(duration)

        while start_time + meeting_duration <= end_time:
            # Check if the committee is already scheduled at this start_time
            if any(s_time <= start_time < e_time for s_time, e_time in scheduled_times[committee]):
                print(f"Committee {committee} is already scheduled at {start_time}. Skipping to next time slot.")
                start_time += timedelta(minutes=15)
                continue

            for room in range(num_rooms):
                print(f"Trying room {room} at {start_time}")
                meeting = {
                    "committee": committee,
                    "duration": duration,
                    "attendees": attendees,
                    "day": day,
                    "room": room,
                    "start_time": start_time,
                    "end_time": start_time + meeting_duration,
                    "conflict_count": 0  # Initialize conflict count
                }

                # Check for conflicts
                if has_conflict(meeting, scheduled_meetings) or has_attendee_conflict(meeting, scheduled_meetings):
                    print(f"Conflict detected for room {room} at {start_time}.")
                else:
                    scheduled_meetings.append(meeting)
                    scheduled_times[committee].append((start_time, start_time + meeting_duration))
                    print(f"Meeting scheduled for committee {committee} in room {room} on {day} at {start_time}")
                    return True

            # Increment start time if all rooms are tried and no suitable room was found
            print(f"All rooms tried and no suitable room found at {start_time}. Trying next time slot.")
            start_time += timedelta(minutes=15)

        print(f"No available time slot found for committee {committee} on {day}")

    print(f"Failed to schedule meeting for committee: {committee}")
    return False






# Schedule meetings
scheduled_meetings = []
for committee, duration in committees:
    attendees = attendees_dict[committee]
    if not schedule_meeting(committee, duration, attendees, scheduled_meetings):
        print(f"Failed to schedule meeting for committee: {committee}")

# Generate CSV file
with open("scheduled_meetings.csv", "w", newline="") as csvfile:
    fieldnames = ["Day", "Start Time", "End Time", "Room", "Committee Code", "Conflict Count"]
    writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
    writer.writeheader()
    for meeting in sorted(scheduled_meetings, key=lambda m: (m["day"], m["start_time"], m["committee"])):
        # Check for scheduling conflicts (this is a simplistic example, adjust according to your logic)
        conflict = "No" if not has_conflict(meeting, scheduled_meetings) and not has_attendee_conflict(meeting, scheduled_meetings) else "Yes"
        writer.writerow(
            {
                "Day": meeting["day"],
                "Start Time": meeting["start_time"].strftime("%Y-%m-%d %H:%M:%S"),
                "End Time": meeting["end_time"].strftime("%Y-%m-%d %H:%M:%S"),
                "Room": meeting["room"],
                "Committee Code": meeting["committee"],
                "Conflict Count": meeting["conflict_count"]
            }
        )


# Generate Word document
doc = Document()
doc.add_heading("AAPM Committee Meeting Schedule")
for day in days.keys():
    doc.add_heading(f"{day.title()} Schedule", level=2)
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Committee Code"
    table.rows[0].cells[1].text = "Duration"
    table.rows[0].cells[2].text = "Room"
    table.rows[0].cells[3].text = "Start Time"
    table.rows[0].cells[4].text = "End Time"
    table.rows[0].cells[5].text = "Chair"
    table.rows[0].cells[6].text = "Attendees"
    for meeting in sorted(
        [m for m in scheduled_meetings if m["day"] == day], key=lambda m: m["start_time"]
    ):
        row = table.add_row()
        row.cells[0].text = meeting["committee"]
        row.cells[1].text = meeting["duration"]
        row.cells[2].text = str(meeting["room"])
        row.cells[3].text = meeting["start_time"].strftime("%H:%M:%S")
        row.cells[4].text = meeting["end_time"].strftime("%H:%M:%S")
        chair_names = [attendee["name"] for attendee in meeting["attendees"] if attendee["ischair"] == "Y"]
        row.cells[5].text = "\n".join(chair_names)
        attendee_names = [attendee["name"] for attendee in meeting["attendees"]]
        row.cells[6].text = "\n".join(attendee_names)
doc.save("committee_schedule.docx")

# Close database connection
cursor.close()
cnxn.close()

